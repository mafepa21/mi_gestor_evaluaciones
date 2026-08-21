#!/usr/bin/env python3
"""Rewrite the active session-sequencing corpus into the weekly CLIL import format.

The source files are intentionally treated as source material, not as a schema. The
rewriter keeps the original detail inside each teacher guide while emitting one
WEEK with a LONG BLOCK and a SHORT BLOCK, each with an importable activity table.
"""

from __future__ import annotations

import argparse
import html
import os
import re
import shutil
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


HEADING_RE = re.compile(r"^(#{2,4})\s+(.+?)\s*$")
WEEK_RE = re.compile(r"^(?:SEMANA|SETMANA|WEEK)\s+(\d+)(?:\s*[.\-–—:]\s*(.*))?$", re.I)
SESSION_RE = re.compile(
    r"^(?:SESI(?:ÓN|ON|ONES|ONS)?|SESSION(?:S)?|SESSI(?:Ó|O)?(?:NS)?)\s+"
    r"(\d+)(?:\s*(?:a|and|to|[-–—])\s*(\d+))?(?:\s*[.\-–—:]\s*(.*))?$",
    re.I,
)
BLOCK_RE = re.compile(r"^(?:BLOQUE|BLOCK)\s+(\d+)(?:\s*[.\-–—:]\s*(.*))?$", re.I)


@dataclass
class Activity:
    time: str
    phase: str
    activity: str
    teacher: str
    students: str
    clil: str
    evidence: str
    materials: str = ""
    adaptations: str = ""
    activity_key: str = ""
    activity_type: str = "core"
    planned_minutes: int = 0
    purpose: str = ""
    organisation: str = ""
    setup: str = ""
    student_instructions: str = ""
    timing_breakdown: str = ""
    slow_group_plan: str = ""
    fast_group_extension: str = ""


@dataclass
class Week:
    number: int
    title: str
    body: str
    long_body: str
    short_body: str
    objective: str
    criteria: str
    materials: str


def clean(value: str) -> str:
    value = re.sub(r"<br\s*/?>", "\n", value, flags=re.I)
    value = re.sub(r"^\s*#{2,6}\s+", "", value, flags=re.M)
    value = re.sub(r"\s+", " ", value.replace("\u00a0", " ")).strip()
    return value.strip("| ")


def plain(value: str) -> str:
    value = re.sub(r"[`*_]+", "", value)
    value = re.sub(r"\[[^]]+\]\([^)]*\)", lambda m: m.group(0).split("]")[0].lstrip("["), value)
    return clean(value)


def norm(value: str) -> str:
    return plain(value).casefold()


TRANSLATIONS = (
    (r"\bcalentamiento\b", "Warm-up"),
    (r"\bparte principal\b", "Main practice"),
    (r"\bvuelta a la calma\b", "Cool-down"),
    (r"\bestiramientos?\b", "Stretching"),
    (r"\bdesplazamientos?\b", "Movement patterns"),
    (r"\bcircuito\b", "Circuit"),
    (r"\bpor parejas\b", "in pairs"),
    (r"\bpor grupos\b", "in groups"),
    (r"\bjuego\b", "game"),
    (r"\btorneo\b", "tournament"),
    (r"\bpartido\b", "match"),
    (r"\breglas?\b", "rules"),
    (r"\bseguridad\b", "safety"),
    (r"\bcooperaci[oó]n\b", "cooperation"),
    (r"\bparticipaci[oó]n\b", "participation"),
    (r"\bobservaci[oó]n docente\b", "teacher observation"),
    (r"\bobservaci[oó]n\b", "observation"),
    (r"\breflexi[oó]n\b", "reflection"),
    (r"\bevaluaci[oó]n\b", "assessment"),
    (r"\bauto[- ]evaluaci[oó]n\b", "self-assessment"),
    (r"\bobjetivo\b", "objective"),
    (r"\bmaterial(?:es)?\b", "materials"),
)


def englishize(value: str) -> str:
    translated = plain(value)
    for pattern, replacement in TRANSLATIONS:
        translated = re.sub(pattern, replacement, translated, flags=re.I)
    return translated


def split_heading_sections(text: str) -> list[tuple[str, str]]:
    lines = text.splitlines()
    starts: list[tuple[int, str]] = []
    for index, line in enumerate(lines):
        match = HEADING_RE.match(line)
        if not match:
            continue
        title = plain(match.group(2))
        lowered = norm(title)
        if WEEK_RE.match(title) or SESSION_RE.match(title) or BLOCK_RE.match(title):
            starts.append((index, title))
        elif re.match(r"^(?:session|sesion|sesión|bloque|block)\s+\d+\b", lowered):
            starts.append((index, title))
    if not starts:
        return [("Sequence", text)]
    sections: list[tuple[str, str]] = []
    for position, (index, title) in enumerate(starts):
        end = starts[position + 1][0] if position + 1 < len(starts) else len(lines)
        sections.append((title, "\n".join(lines[index + 1 : end]).strip()))
    return sections


def number_title(title: str, fallback: int) -> tuple[int, str]:
    if match := WEEK_RE.match(title):
        return int(match.group(1)), clean(match.group(2) or "") or title
    if match := BLOCK_RE.match(title):
        return int(match.group(1)), clean(match.group(2) or "") or title
    if match := SESSION_RE.match(title):
        return int(match.group(1)), clean(match.group(3) or "") or title
    return fallback, title


def expand_ranges(sections: list[tuple[str, str]]) -> list[tuple[int, str, str]]:
    expanded: list[tuple[int, str, str]] = []
    for fallback, (title, body) in enumerate(sections, 1):
        number, label = number_title(title, fallback)
        match = SESSION_RE.match(title)
        if match and match.group(2):
            end = int(match.group(2))
            for current in range(number, end + 1):
                expanded.append((current, f"Session {current} - {label}", body))
        else:
            expanded.append((number, label, body))
    expanded.sort(key=lambda item: item[0])
    return expanded


def value_after_label(body: str, labels: Iterable[str]) -> str:
    label_pattern = "|".join(re.escape(label) for label in labels)
    lines = body.splitlines()
    for index, line in enumerate(lines):
        match = re.match(rf"^\s*(?:\*\*)?(?:{label_pattern})(?:\*\*)?\s*[:\-]\s*(.*)$", line, re.I)
        if match:
            value = englishize(match.group(1))
            if value:
                return value
            for following in lines[index + 1 :]:
                candidate = englishize(following)
                if candidate:
                    return candidate
            continue
        if norm(line).strip("#* ") in {norm(label) for label in labels}:
            for following in lines[index + 1 :]:
                candidate = englishize(following)
                if candidate and not candidate.startswith("#"):
                    return candidate.lstrip("- ")
    return ""


def split_variants(body: str) -> tuple[str, str, str]:
    lines = body.splitlines()
    markers: list[tuple[int, str]] = []
    for index, line in enumerate(lines):
        marker_text = re.sub(r"^[#*\s]+", "", line).strip()
        lowered = norm(marker_text)
        if re.match(r"(?:long block|bloque largo|double version|versión doble|version doble)\b", lowered):
            markers.append((index, "long"))
        elif re.match(r"(?:short block|bloque corto|simple version|versión simple|version simple)\b", lowered):
            markers.append((index, "short"))
    if not markers:
        return body, body, body
    shared = "\n".join(lines[: markers[0][0]]).strip()
    long_parts: list[str] = []
    short_parts: list[str] = []
    for position, (index, kind) in enumerate(markers):
        end = markers[position + 1][0] if position + 1 < len(markers) else len(lines)
        content = "\n".join(lines[index + 1 : end]).strip()
        if kind == "long":
            long_parts.append(content)
        else:
            short_parts.append(content)
    long_body = "\n".join(part for part in long_parts if part).strip() or body
    short_body = "\n".join(part for part in short_parts if part).strip() or body
    return shared, long_body, short_body


def table_rows(body: str) -> list[Activity]:
    lines = body.splitlines()
    activities: list[Activity] = []
    index = 0
    while index < len(lines):
        if not lines[index].lstrip().startswith("|"):
            index += 1
            continue
        table: list[list[str]] = []
        while index < len(lines) and lines[index].lstrip().startswith("|"):
            cells = [clean(cell) for cell in lines[index].strip().strip("|").split("|")]
            if cells and not all(re.fullmatch(r"[-: ]+", cell or "-") for cell in cells):
                table.append(cells)
            index += 1
        if len(table) < 2:
            continue
        header = [norm(cell) for cell in table[0]]

        def column(tokens: tuple[str, ...]) -> int | None:
            for position, cell in enumerate(header):
                if any(token in ("activity", "actividad") for token in tokens) and ("activity id" in cell or "activity key" in cell):
                    continue
                if any(token in cell for token in tokens):
                    return position
            return None

        time_col = column(("time", "hora", "tiempo"))
        id_col = column(("activity id", "activity key", "id actividad", "clave actividad"))
        type_col = column(("type", "activity type", "tipo"))
        minutes_col = column(("minutes", "minutos", "duration", "duracion"))
        phase_col = column(("phase", "fase"))
        activity_col = next((position for position, cell in enumerate(header) if cell in {"activity", "actividad", "task", "tarea"}), None)
        if activity_col is None:
            activity_col = column(("activity", "actividad", "task", "tarea"))
        purpose_col = column(("purpose", "proposito"))
        organisation_col = column(("organisation", "organization", "organizacion", "grouping"))
        setup_col = column(("set-up", "setup", "preparation", "preparacion"))
        teacher_col = column(("teacher", "profesor", "docente"))
        student_output_col = column(("student output", "student instructions", "output alumno"))
        student_col = column(("student actions", "student role", "student does", "student", "alumno", "learner", "alumnado"))
        timing_col = column(("timing breakdown", "desglose temporal", "timing"))
        clil_col = column(("clil", "language", "lengua", "scaffolding", "andamiaje"))
        evidence_col = column(("evidence", "evidencia", "assessment", "evaluacion"))
        materials_col = column(("material", "materials", "materiales"))
        adaptation_col = column(("adaptation", "adaptaciones", "inclusion"))
        slow_col = column(("slow group", "if the group is slow", "grupo lento"))
        fast_col = column(("fast group", "if the group is ahead", "grupo adelantado", "extension"))

        def get(row: list[str], position: int | None) -> str:
            return row[position] if position is not None and position < len(row) else ""

        for row in table[1:]:
            activity = get(row, activity_col)
            if not activity:
                activity = " / ".join(value for value in row if value)
            if not activity:
                continue
            # A previously generated file can contain metadata tables or rows copied from the
            # teacher sheet. They are not executable activities and must never enter QUICK VIEW.
            if norm(activity).startswith(("objective", "objectives", "criteria", "criterion", "materials", "saberes", "assessment", "evidence focus", "long block", "short block")):
                continue
            activities.append(
                Activity(
                    time=get(row, time_col),
                    phase=get(row, phase_col),
                    activity=plain(activity),
                    teacher=plain(get(row, teacher_col)),
                    students=plain(get(row, student_col)),
                    clil=plain(get(row, clil_col)),
                    evidence=plain(get(row, evidence_col)),
                    materials=plain(get(row, materials_col)),
                    adaptations=plain(get(row, adaptation_col)),
                    activity_key=plain(get(row, id_col)),
                    activity_type=plain(get(row, type_col)) or "core",
                    planned_minutes=int(re.search(r"\d+", get(row, minutes_col)).group(0)) if re.search(r"\d+", get(row, minutes_col)) else 0,
                    purpose=plain(get(row, purpose_col)),
                    organisation=plain(get(row, organisation_col)),
                    setup=plain(get(row, setup_col)),
                    student_instructions=plain(get(row, student_output_col)),
                    timing_breakdown=plain(get(row, timing_col)),
                    slow_group_plan=plain(get(row, slow_col)),
                    fast_group_extension=plain(get(row, fast_col)),
                )
            )
    return activities


def chunks_from_text(body: str) -> list[Activity]:
    lines = body.splitlines()
    candidates: list[tuple[str, str]] = []
    current_title = "Teaching sequence"
    current: list[str] = []
    phase_tokens = (
        "warm-up", "warm up", "calentamiento", "briefing", "asamblea", "main part",
        "parte principal", "practice", "práctica", "station", "estacion", "closure",
        "cierre", "cool-down", "cool down", "vuelta a la calma", "evaluation", "evaluación",
        "reflection", "reflexión", "challenge", "reto", "tournament", "torneo",
    )

    def flush() -> None:
        nonlocal current
        text = plain(" ".join(current))
        if text:
            candidates.append((current_title, text))
        current = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if re.match(r"^\s*(?:\*\*)?(?:objective|objectives|objetivo|objetivos|criteria|criterion|criterios|criterio|materials|material|materiales|evaluation|evaluación)\b", stripped, re.I):
            continue
        heading = HEADING_RE.match(stripped)
        title = plain(heading.group(2)) if heading else ""
        lowered = norm(title or stripped)
        is_phase = bool(title) and any(token in lowered for token in phase_tokens)
        time_match = re.match(r"(?:⏱️\s*)?([0-9]{1,3}\s*['′]?(?:\s*[-–—]\s*[0-9]{1,3}\s*['′]?)?)\s*:?\s*(.*)", stripped)
        if is_phase or time_match:
            flush()
            current_title = englishize(title or (time_match.group(2) if time_match else stripped))
            if time_match and time_match.group(2):
                current_title = englishize(time_match.group(2))
            continue
        if heading:
            flush()
            current_title = "Practice"
            continue
        if stripped.startswith("-") or re.match(r"^\d+[.)]\s", stripped):
            if len(current) > 2:
                flush()
            current.append(stripped.lstrip("- ").lstrip("0123456789.) "))
        else:
            current.append(stripped)
    flush()
    if not candidates:
        text = plain(body)
        candidates = [("Teaching sequence", text[:1800] or "Follow the planned sequence.")]

    activities: list[Activity] = []
    for title, text in candidates:
        name = englishize(title)
        bold_match = re.search(r"\*\*([^*]{3,100})\*\*", text)
        if bold_match:
            name = englishize(bold_match.group(1))
        lowered = norm(title)
        if any(token in lowered for token in ("warm", "calent", "briefing", "asamblea")):
            students = "Follow the signal, maintain safe spacing, and name the key movement or rule."
        elif any(token in lowered for token in ("closure", "cierre", "cool", "vuelta", "evaluation", "evaluación", "reflection", "reflexión")):
            students = "Recover, record the evidence, and state one improvement or next step."
        else:
            students = "Perform the task, rotate roles, and explain one decision or adjustment to a partner."
        activities.append(
            Activity(
                time="",
                phase=englishize(title),
                activity=name,
                teacher=text[:1800],
                students=students,
                clil="",
                evidence="",
            )
        )
    return activities


def enrich(activities: list[Activity], subject: str) -> list[Activity]:
    subject_words = "pass, receive, defend, space, support" if any(
        word in norm(subject) for word in ("basket", "handball", "balonmano", "volley", "voleibol", "colpbol", "frisbee", "pilota")
    ) else "warm-up, effort, posture, recovery, feedback"
    result: list[Activity] = []
    for item in activities:
        clil = englishize(item.clil) or (
            f"Language goal: explain a choice and give respectful feedback. Vocabulary: {subject_words}. "
            "Sentence stems: “We choose… because…”, “I noticed…”, “Could you show…?”"
        )
        evidence = englishize(item.evidence) or "Teacher checklist: safe execution, active participation, target strategy, and one spoken English contribution."
        source_detail = englishize(item.teacher)
        teacher = item.teacher if item.teacher.startswith("Teacher cue (English):") else (
            "Teacher cue (English): state the aim and success criteria; model the first step; check safety and understanding; "
            "then circulate and give one precise feedback cue. "
            f"Lesson detail: {source_detail}"
            if source_detail
            else "Teacher cue (English): state the aim and success criteria; model the first step; check safety and understanding; then circulate and give one precise feedback cue."
        )
        students = englishize(item.students) or "Complete the task, rotate roles, and report one observation using the sentence stem."
        student_instructions = englishize(item.student_instructions) or students
        materials = englishize(item.materials) or "Equipment and visual cards named in the lesson setup; observation or recording sheet."
        adaptations = englishize(item.adaptations) or "Use a visual model, adjust space, distance or equipment, and keep the evidence target stable."
        purpose = englishize(item.purpose) or f"Practise {englishize(item.activity).lower()} safely and produce the evidence named below."
        organisation = englishize(item.organisation) or "Pairs or groups of four; assign performer, coach, observer and equipment roles before starting."
        setup = englishize(item.setup) or "Place equipment and display the task card before the timer starts; keep one clear route for transitions."
        timing_breakdown = englishize(item.timing_breakdown) or "Includes a short teacher briefing, active practice, one role change, and a visible stop signal."
        slow_group = englishize(item.slow_group_plan) or "Drop the optional variation first, keep one clear success criterion, and use a visible countdown for the transition."
        fast_group = englishize(item.fast_group_extension) or "Repeat once with a role swap or add one controlled decision; do not add a new setup."
        result.append(Activity(
            time=item.time,
            phase=englishize(item.phase) or "Practice",
            activity=englishize(item.activity),
            teacher=teacher,
            students=students,
            clil=clil,
            evidence=evidence,
            materials=materials,
            adaptations=adaptations,
            activity_key=item.activity_key,
            activity_type=item.activity_type or "core",
            planned_minutes=item.planned_minutes,
            purpose=purpose,
            organisation=organisation,
            setup=setup,
            student_instructions=student_instructions,
            timing_breakdown=timing_breakdown,
            slow_group_plan=slow_group,
            fast_group_extension=fast_group,
        ))
    return result


def interval_minutes(value: str) -> int:
    match = re.search(r"(\d{1,3})\s*[′'’]?\s*[-–—]\s*(\d{1,3})", value or "")
    if not match:
        return 0
    start, end = int(match.group(1)), int(match.group(2))
    return max(0, end - start)


def classify_activity(item: Activity) -> str:
    value = norm(f"{item.phase} {item.activity}")
    if any(token in value for token in ("entry", "entrada", "briefing", "organisation", "organizacion", "setup")):
        return "setup"
    if any(token in value for token in ("closure", "cierre", "cool-down", "cool down", "vuelta a la calma", "exit", "reflection", "reflexion")):
        return "closure"
    if item.activity_type and item.activity_type.casefold() not in {"core", "practice"}:
        return item.activity_type
    return "core"


def consolidate_for_block(activities: list[Activity], total: int) -> list[Activity]:
    """Keep the operating card realistic for a class of 35 without deleting teacher intent."""
    if len(activities) <= 1:
        return activities
    maximum = 4 if total >= 60 else 3
    if len(activities) <= maximum:
        return activities
    head = activities[: maximum - 1]
    tail = activities[maximum - 1 :]
    merged = Activity(
        time="",
        phase="Consolidation and closure",
        activity="Consolidated practice, evidence check and exit response",
        teacher="Use one visible success criterion. Give the final signal, check the evidence quickly in pairs, and collect only the agreed exit response.\n\n" + "\n".join(item.teacher for item in tail),
        students="Complete the essential final task, check the partner evidence, and submit one concise improvement statement.",
        clil="Language goal: summarise evidence and state a next step. Sentence stem: “Our evidence shows…; next we will…”.",
        evidence="One completed exit response plus the partner/teacher evidence already collected.",
        materials="Use the materials already in place; no new station or distribution is introduced.",
        adaptations="If time is short, use the written exit response instead of plenary sharing.",
        activity_type="closure",
        purpose="Close the lesson without adding another setup or a full-class reporting round.",
        organisation="Pairs first; teacher samples a small number of responses instead of hearing 35 students one by one.",
        setup="Keep students in their current groups and place the exit prompt where everyone can see it.",
        student_instructions="Compare, record, and hand in the essential evidence.",
        timing_breakdown="Includes the final role change, evidence check and exit response; no whole-class round is assumed.",
        slow_group_plan="Use the written exit response and defer optional sharing.",
        fast_group_extension="Add one peer feedback sentence, not a new activity.",
    )
    return head + [merged]


def assign_times(activities: list[Activity], total: int) -> list[Activity]:
    """Assign conservative windows with setup, transitions and closure included.

    Source times are treated as evidence of pedagogical order, not as a promise that 35
    students can change stations instantly. The result always fits the selected block.
    """
    activities = consolidate_for_block(activities, total)
    if not activities:
        return []
    count = len(activities)
    transition = 2 if total >= 60 else 1
    setup = 5 if total >= 60 else 4
    closure = 5 if total >= 60 else 3
    if count == 1:
        setup, transition, closure = 3, 0, 3
    active_budget = max(count, total - setup - closure - transition * max(0, count - 1))
    source_durations = [interval_minutes(item.time) or item.planned_minutes for item in activities]
    if not any(source_durations):
        source_durations = [1] * count
    scale = active_budget / max(1, sum(source_durations))
    durations = [max(1, round(value * scale)) for value in source_durations]
    difference = active_budget - sum(durations)
    durations[-1] = max(1, durations[-1] + difference)
    result: list[Activity] = []
    cursor = 0
    for index, (item, duration) in enumerate(zip(activities, durations)):
        if index == 0:
            cursor += setup
        start = cursor
        end = min(total - closure if index < count - 1 else total, start + duration)
        if index == count - 1:
            end = total
        if end <= start:
            end = min(total, start + 1)
        activity_type = classify_activity(item)
        label = f"{start}′–{end}′"
        minutes = end - start
        timing = item.timing_breakdown or f"{minutes} min total: briefing/model, active practice, role change and evidence check."
        result.append(Activity(
            time=label,
            phase=item.phase,
            activity=item.activity,
            teacher=item.teacher,
            students=item.students,
            clil=item.clil,
            evidence=item.evidence,
            materials=item.materials,
            adaptations=item.adaptations,
            activity_key=item.activity_key,
            activity_type=activity_type,
            planned_minutes=minutes,
            purpose=item.purpose,
            organisation=item.organisation,
            setup=item.setup,
            student_instructions=item.student_instructions,
            timing_breakdown=timing,
            slow_group_plan=item.slow_group_plan,
            fast_group_extension=item.fast_group_extension,
        ))
        cursor = end + (transition if index < count - 1 else 0)
    return result


def build_block_activities(week: Week, body: str, total: int, block_code: str, subject: str) -> list[Activity]:
    activities = table_rows(body) or chunks_from_text(body)
    activities = enrich(activities, subject)
    activities = assign_times(activities, total)
    result: list[Activity] = []
    for index, item in enumerate(activities, 1):
        result.append(Activity(
            time=item.time,
            phase=item.phase,
            activity=item.activity,
            teacher=item.teacher,
            students=item.students,
            clil=item.clil,
            evidence=item.evidence,
            materials=item.materials,
            adaptations=item.adaptations,
            activity_key=f"W{week.number:02d}-{block_code}-{index:02d}",
            activity_type=item.activity_type,
            planned_minutes=item.planned_minutes,
            purpose=item.purpose,
            organisation=item.organisation,
            setup=item.setup,
            student_instructions=item.student_instructions,
            timing_breakdown=item.timing_breakdown,
            slow_group_plan=item.slow_group_plan,
            fast_group_extension=item.fast_group_extension,
        ))
    return result


def make_weeks(text: str) -> list[Week]:
    title = plain(next((line[2:].strip() for line in text.splitlines() if line.startswith("# ")), "Session sequence"))
    sections = split_heading_sections(text)
    # Building Health has an archived, pedagogically richer ten-session source. Pair the
    # double/simple sessions into five weekly alternatives so the app can place the long
    # block on either day without duplicating the short activity list.
    if any(SESSION_RE.match(section_title) and re.search(r"\b(?:double|simple)\b", norm(section_title)) for section_title, _ in sections):
        paired: dict[int, dict[str, tuple[str, str]]] = {}
        for section_title, section_body in sections:
            match = SESSION_RE.match(section_title)
            if not match:
                continue
            numbers = [int(match.group(1))]
            if match.group(2):
                numbers.append(int(match.group(2)))
            variant = "long" if "double" in norm(section_title) else "short"
            for number in numbers:
                week_number = (number + 1) // 2
                paired.setdefault(week_number, {})[variant] = (section_title, section_body)
        if paired:
            weeks: list[Week] = []
            for week_number, variants in sorted(paired.items()):
                long_title, long_body = variants.get("long", (f"Week {week_number} long", ""))
                short_title, short_body = variants.get("short", (f"Week {week_number} short", ""))
                source_body = long_body or short_body
                objective = value_after_label(source_body, ("Specific Objective", "Objective", "Objectives"))
                criteria = value_after_label(source_body, ("Evaluation Criteria Worked", "Criteria", "Criterion"))
                materials = value_after_label(source_body, ("Required Materials", "Materials", "Material"))
                weeks.append(Week(
                    week_number,
                    plain(long_title.split(".", 1)[-1] if "." in long_title else long_title),
                    "",
                    long_body,
                    short_body,
                    objective,
                    criteria,
                    materials,
                ))
            return weeks
    if "**Import rule:**" in text:
        canonical_sections: list[tuple[str, str]] = []
        for section_title, section_body in sections:
            if not WEEK_RE.match(section_title):
                continue
            _, section_label = number_title(section_title, 1)
            if re.match(r"(?:long block|short block)\b", norm(section_label)):
                continue
            canonical_sections.append((section_title, section_body))
        if canonical_sections:
            sections = canonical_sections
    expanded = expand_ranges(sections)
    weeks: list[Week] = []
    for number, label, body in expanded:
        shared, long_body, short_body = split_variants(body)
        objective = value_after_label(shared + "\n" + body, ("Objective", "Objectives", "Objetivo", "Objetivos"))
        criteria = value_after_label(shared + "\n" + body, ("Criteria", "Criterion", "Criterios", "Criterio", "Evaluation"))
        materials = value_after_label(shared + "\n" + body, ("Materials", "Material", "Materiales"))
        weeks.append(Week(number, label or f"Week {number}", shared, long_body, short_body, objective, criteria, materials))
    if not weeks:
        weeks.append(Week(1, title, text, text, text, "", "", ""))
    return weeks


def markdown_table(activities: list[Activity]) -> str:
    rows = [
        "| Time | Activity ID | Type | Minutes | Phase | Activity | Organisation | Student output | Materials | Evidence |",
        "|---|---|---|---:|---|---|---|---|---|---|",
    ]
    for item in activities:
        cells = [item.time, item.activity_key, item.activity_type, str(item.planned_minutes), item.phase, item.activity, item.organisation, item.student_instructions, item.materials, item.evidence]
        rows.append("| " + " | ".join(clean(cell).replace("|", "/").replace("\n", "<br>") for cell in cells) + " |")
    return "\n".join(rows)


def activity_details_markdown(activities: list[Activity]) -> str:
    lines = ["### ACTIVITY DETAILS", "", "Each Activity ID below matches one row in QUICK VIEW. Use this section when the teacher needs the full operating instructions.", ""]
    labels = (
        ("Purpose", "purpose"),
        ("Organisation", "organisation"),
        ("Set-up", "setup"),
        ("Teacher instructions", "teacher"),
        ("Instructions for students", "student_instructions"),
        ("Timing breakdown", "timing_breakdown"),
        ("CLIL focus", "clil"),
        ("Materials", "materials"),
        ("Evidence", "evidence"),
        ("Adaptations", "adaptations"),
        ("If the group is slow", "slow_group_plan"),
        ("If the group is ahead", "fast_group_extension"),
    )
    for item in activities:
        lines.append(f"#### ACTIVITY {item.activity_key}")
        lines.append("")
        for label, field in labels:
            value = clean(getattr(item, field))
            if value:
                lines.append(f"**{label}:** {value}")
                lines.append("")
    return "\n".join(lines).rstrip()


def render_markdown(source: Path, title: str, weeks: list[Week]) -> str:
    lines = [
        f"# {title} - Weekly Session Sequence",
        "",
        "**Language:** English delivery with CLIL support and explicit teacher prompts.",
        "**Planning rule:** Each week contains one LONG BLOCK (90 effective minutes) and one SHORT BLOCK (30 effective minutes). The app assigns the long block to two consecutive timetable slots and the short block to one simple slot, independently of whether the group has the long day on Monday or Friday.",
        "**Import rule:** The `QUICK VIEW` table is the structured contract used by the app. Every row has a stable Activity ID; the `ACTIVITY DETAILS` section expands the same ID and is never imported as a duplicate activity.",
        "",
        "## Teacher operating key",
        "",
        "- **Before class:** prepare the listed materials, display the CLIL sentence stems, and assign roles before movement starts.",
        "- **During class:** use the teacher instructions as cues, observe the evidence column, and adapt space, distance, equipment, or role without removing participation.",
        "- **After class:** record one piece of evidence and one next step in the session journal.",
        "",
        f"**Source file normalised:** `{source.name}`. The pedagogical source detail is retained in each activity teacher guide and in the final trace section.",
        "",
    ]
    for week in weeks:
        shared_activities: list[Activity] = []
        long_activities = build_block_activities(week, week.long_body, 90, "L", title)
        short_activities = build_block_activities(week, week.short_body, 30, "S", title)
        lines += [
            f"## WEEK {week.number} - {plain(week.title)}",
            "",
            f"**Objective:** {englishize(week.objective) or 'Achieve the shared learning intention through safe participation, purposeful practice and reflection.'}",
            f"**Criteria / evidence focus:** {englishize(week.criteria) or 'Safe execution, cooperation, decision-making, communication and reflective improvement.'}",
            f"**Materials:** {englishize(week.materials) or 'Use the equipment, visual cards and recording sheet named in the activity table.'}",
            "",
            "### LONG BLOCK (90 effective minutes)",
            "",
            "#### QUICK VIEW",
            "",
            markdown_table(long_activities),
            "",
            activity_details_markdown(long_activities),
            "",
            "### SHORT BLOCK (30 effective minutes)",
            "",
            "#### QUICK VIEW",
            "",
            markdown_table(short_activities),
            "",
            activity_details_markdown(short_activities),
            "",
            "### Inclusion and teacher adaptations",
            "",
            "- Offer a visual model, reduced space or distance, lighter or larger equipment, peer support, and an opt-in intensity level.",
            "- Keep the evidence target stable while changing the route to success; never use elimination as the default adaptation.",
            "- Check comprehension with “Show me the first step” before increasing speed, opposition or complexity.",
            "",
        ]
    lines += [
        "## Source trace",
        "",
        "The source sequence was reorganised into a weekly long/short contract. Original activity wording remains embedded in the teacher-guide cells so the teacher can cross-check the planning without reopening a second document.",
        "",
    ]
    return "\n".join(lines).rstrip() + "\n"


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    width = properties.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        properties.append(width)
    width.set(qn("w:w"), str(width_dxa))
    width.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    properties = tbl.tblPr
    width = properties.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:w"), str(sum(widths)))
    width.set(qn("w:type"), "dxa")
    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width_dxa in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width_dxa))
        grid.append(column)
    for row in table.rows:
        for cell, width_dxa in zip(row.cells, widths):
            set_cell_width(cell, width_dxa)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def set_run_font(run, name: str = "Calibri", size: float = 10, color: str = "1F2937", bold: bool = False) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def add_text(doc: Document, text: str, style: str = "Normal", size: float = 10, color: str = "1F2937", bold: bool = False) -> None:
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def add_activity_table(doc: Document, activities: list[Activity]) -> None:
    headers = ["Time", "Phase", "Activity", "Teacher instructions", "Student actions", "CLIL focus", "Evidence", "Materials", "Adaptations"]
    table = doc.add_table(rows=1, cols=len(headers))
    widths = [650, 800, 1650, 2150, 1550, 1500, 1250, 1250, 1200]
    set_table_geometry(table, widths)


def add_quick_view_table(doc: Document, activities: list[Activity]) -> None:
    headers = ["Time", "Activity ID", "Type", "Minutes", "Phase", "Activity", "Organisation", "Student output", "Materials", "Evidence"]
    table = doc.add_table(rows=1, cols=len(headers))
    widths = [700, 950, 700, 500, 850, 2100, 1800, 1800, 1500, 1450]
    set_table_geometry(table, widths)
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, "DCE9F7")
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        set_run_font(run, size=7.2, color="0B2545", bold=True)
    header_row = table.rows[0]._tr
    tr_properties = header_row.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_properties.append(repeat)
    for item in activities:
        row = table.add_row()
        row_properties = row._tr.get_or_add_trPr()
        row_properties.append(OxmlElement("w:cantSplit"))
        values = [item.time, item.activity_key, item.activity_type, str(item.planned_minutes), item.phase, item.activity, item.organisation, item.student_instructions, item.materials, item.evidence]
        for cell, value in zip(row.cells, values):
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            run = paragraph.add_run(value)
            set_run_font(run, size=7.0, color="1F2937")
    set_table_geometry(table, widths)


def add_activity_detail_cards(doc: Document, activities: list[Activity]) -> None:
    labels = (
        ("Purpose", "purpose"),
        ("Organisation", "organisation"),
        ("Set-up", "setup"),
        ("Teacher instructions", "teacher"),
        ("Instructions for students", "student_instructions"),
        ("Timing breakdown", "timing_breakdown"),
        ("CLIL focus", "clil"),
        ("Materials", "materials"),
        ("Evidence", "evidence"),
        ("Adaptations", "adaptations"),
        ("If the group is slow", "slow_group_plan"),
        ("If the group is ahead", "fast_group_extension"),
    )
    add_text(doc, "Activity details", style="Heading 3", size=12, color="1F4D78", bold=True)
    for item in activities:
        add_text(doc, f"ACTIVITY {item.activity_key}", size=10, color="0B2545", bold=True)
        for label, field in labels:
            value = getattr(item, field)
            if not value:
                continue
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.15)
            paragraph.paragraph_format.space_after = Pt(3)
            run = paragraph.add_run(f"{label}: ")
            set_run_font(run, size=8.2, color="4B5563", bold=True)
            run = paragraph.add_run(value)
            set_run_font(run, size=8.2, color="1F2937")
def write_docx(path: Path, title: str, weeks: list[Week], source: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.15
    for style_name, size, color, before, after in (("Heading 1", 16, "2E74B5", 18, 10), ("Heading 2", 13, "2E74B5", 14, 7), ("Heading 3", 12, "1F4D78", 10, 5)):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run(f"Session sequencing | {title[:70]}")
    set_run_font(run, size=8.5, color="6B7280")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Teacher working document | CLIL session card")
    set_run_font(run, size=8, color="6B7280")

    add_text(document, title, size=22, color="0B2545", bold=True)
    add_text(document, "Weekly session sequence | English delivery with CLIL support", size=11.5, color="4B5563")
    add_text(document, "Planning rule: one long block (90 effective minutes) plus one short block (30 effective minutes) per week. The app maps the long block to two consecutive timetable slots and the short block to one simple slot.", size=9.5, color="374151")
    add_text(document, f"Source normalised: {source.name}", size=8.5, color="6B7280")
    add_text(document, "Teacher operating key", style="Heading 1", size=16, color="2E74B5", bold=True)
    for text in (
        "Before class: prepare the materials, display the CLIL sentence stems, and assign roles.",
        "During class: use the teacher instructions as cues, observe the evidence column, and adapt the route to success.",
        "After class: record one piece of evidence and one next step in the session journal.",
    ):
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(text)
        set_run_font(run, size=9.5)

    for index, week in enumerate(weeks):
        if index:
            document.add_page_break()
        long_activities = build_block_activities(week, week.long_body, 90, "L", title)
        short_activities = build_block_activities(week, week.short_body, 30, "S", title)
        add_text(document, f"WEEK {week.number} - {plain(week.title)}", style="Heading 1", size=16, color="2E74B5", bold=True)
        add_text(document, f"Objective: {englishize(week.objective) or 'Achieve the shared learning intention through safe participation, purposeful practice and reflection.'}", size=9.5)
        add_text(document, f"Criteria / evidence focus: {englishize(week.criteria) or 'Safe execution, cooperation, decision-making, communication and reflective improvement.'}", size=9.5)
        add_text(document, f"Materials: {englishize(week.materials) or 'Use the equipment, visual cards and recording sheet named in the activity table.'}", size=9.5)
        add_text(document, "LONG BLOCK | 90 effective minutes", style="Heading 2", size=13, color="2E74B5", bold=True)
        add_text(document, "Quick view", style="Heading 3", size=11.5, color="1F4D78", bold=True)
        add_quick_view_table(document, long_activities)
        add_activity_detail_cards(document, long_activities)
        add_text(document, "SHORT BLOCK | 30 effective minutes", style="Heading 2", size=13, color="2E74B5", bold=True)
        add_text(document, "Quick view", style="Heading 3", size=11.5, color="1F4D78", bold=True)
        add_quick_view_table(document, short_activities)
        add_activity_detail_cards(document, short_activities)
        add_text(document, "Inclusion and teacher adaptations", style="Heading 3", size=12, color="1F4D78", bold=True)
        add_text(document, "Offer a visual model, reduced space or distance, lighter or larger equipment, peer support, and an opt-in intensity level. Keep the evidence target stable while changing the route to success.", size=9)

    atomic_write_docx(path, document)


def atomic_write_text(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary_path: Path | None = None
    try:
        with tempfile.NamedTemporaryFile(
            mode="w",
            encoding="utf-8",
            dir=path.parent,
            prefix=f".{path.name}.",
            suffix=".tmp",
            delete=False,
        ) as handle:
            temporary_path = Path(handle.name)
            handle.write(content)
            handle.flush()
            os.fsync(handle.fileno())
        os.replace(temporary_path, path)
        temporary_path = None
    finally:
        if temporary_path is not None:
            temporary_path.unlink(missing_ok=True)


def atomic_write_docx(path: Path, document: Document) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary_path: Path | None = None
    try:
        with tempfile.NamedTemporaryFile(
            dir=path.parent,
            prefix=f".{path.name}.",
            suffix=".tmp",
            delete=False,
        ) as handle:
            temporary_path = Path(handle.name)
        document.save(temporary_path)
        os.replace(temporary_path, path)
        temporary_path = None
    finally:
        if temporary_path is not None:
            temporary_path.unlink(missing_ok=True)


def backup_file(path: Path, root: Path, backup_dir: Path) -> None:
    if not path.exists():
        return
    relative_path = path.relative_to(root)
    destination = backup_dir / relative_path
    destination.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(path, destination)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--root",
        type=Path,
        required=True,
        help="Root folder containing the Situaciones de aprendizaje corpus.",
    )
    parser.add_argument("--start", type=int, default=0)
    parser.add_argument("--limit", type=int, default=0)
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Parse and render the corpus without changing Markdown or DOCX files.",
    )
    parser.add_argument(
        "--backup-dir",
        type=Path,
        help="Copy each existing source and DOCX here before replacing it.",
    )
    args = parser.parse_args()
    root = args.root.expanduser().resolve()
    if not root.is_dir():
        parser.error(f"--root is not a directory: {root}")
    if not any(path.is_dir() for path in root.rglob("02_SESIONES")):
        parser.error(f"--root does not contain a 02_SESIONES corpus: {root}")
    backup_dir = args.backup_dir.expanduser().resolve() if args.backup_dir else None
    sources = sorted(path for path in root.rglob("*.md") if path.parent.name == "02_SESIONES" and "99_FUENTES" not in path.parts)
    if args.start:
        sources = sources[args.start :]
    if args.limit:
        sources = sources[: args.limit]
    for source in sources:
        archived_source = source.parent.parent / "99_FUENTES" / "version_en_ingles_archivada" / "02_SESIONES" / source.name
        text = archived_source.read_text(encoding="utf-8") if archived_source.exists() else source.read_text(encoding="utf-8")
        title = plain(next((line[2:].strip() for line in text.splitlines() if line.startswith("# ")), source.parent.parent.name))
        title = re.sub(r"(?:\s*-\s*Weekly Session Sequence)+\s*$", "", title, flags=re.I)
        weeks = make_weeks(text)
        markdown = render_markdown(source, title, weeks)
        docx_path = source.with_suffix(".docx")
        if not args.dry_run:
            if backup_dir:
                backup_file(source, root, backup_dir)
                backup_file(docx_path, root, backup_dir)
            atomic_write_text(source, markdown)
            write_docx(docx_path, title, weeks, source)
    action = "would_rewrite" if args.dry_run else "rewritten"
    print(f"{action}={len(sources)}")


if __name__ == "__main__":
    main()
